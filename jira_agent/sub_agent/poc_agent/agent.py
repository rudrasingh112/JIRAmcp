from google.adk.agents import Agent
from google.adk.tools import google_search_tool
from .prompt import poc_agent_instruction
from google.genai import types
from google.genai.types import HttpRetryOptions
import os
import re
from docx import Document
from docx.shared import RGBColor
from fpdf import FPDF


def _add_formatted_runs(para, text: str):
    """Add text to a paragraph, converting **...**  markers to bold runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part:
            para.add_run(part)


def generate_document(content: str, filename: str = "output", filetype: str = ".word"):
    """
    Generates a formatted document from the provided content.

    Supports inline **bold** markers, lines that are entirely **heading**,
    lines starting with '- ' or '* ' for bullet points, and lines starting
    with '#' / '##' for headings.  The very first non-empty line is rendered
    as yellow bold (company name).

    Args:
        content: The text content (may contain **bold** markdown markers).
        filename: The desired name of the file (defaults to 'output').
        filetype: Either '.word' or '.pdf' (defaults to '.word').
    """
    base_name = os.path.splitext(filename)[0]

    if filetype.lower() == ".word":
        full_path = f"{base_name}.docx"
        doc = Document()
        first_content_line_done = False

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            # First non-empty line → company name: yellow + bold
            if not first_content_line_done:
                para = doc.add_paragraph()
                run = para.add_run(re.sub(r'\*+', '', stripped))
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0x00)  # yellow
                first_content_line_done = True
                continue

            # Entire line is a bold heading: **Heading Text**
            heading_match = re.match(r'^\*\*(.+)\*\*$', stripped)
            if heading_match:
                para = doc.add_paragraph()
                run = para.add_run(heading_match.group(1))
                run.bold = True

            # Markdown-style headings: # or ##
            elif re.match(r'^#{1,2} ', stripped):
                para = doc.add_paragraph()
                run = para.add_run(re.sub(r'^#{1,2} ', '', stripped))
                run.bold = True

            # Bullet points
            elif stripped.startswith('- ') or stripped.startswith('* '):
                para = doc.add_paragraph(style='List Bullet')
                _add_formatted_runs(para, stripped[2:])

            # Regular paragraph (may contain inline **bold**)
            else:
                para = doc.add_paragraph()
                _add_formatted_runs(para, stripped)

        doc.save(full_path)

    elif filetype.lower() == ".pdf":
        full_path = f"{base_name}.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=content)
        pdf.output(full_path)

    else:
        return f"Error: Unsupported filetype '{filetype}'. Use '.word' or '.pdf'."

    return f"Successfully generated {full_path}"

poc_agent = Agent(
    model='gemini-2.5-pro',
    name='poc_agent',
    description='A clear and descriptive document generation agent.',
    instruction=poc_agent_instruction,
    tools = [generate_document],
    generate_content_config= types.GenerateContentConfig(
        temperature=0.7,
        max_output_tokens=50000,
        http_options=types.HttpOptions(
            retry_options=HttpRetryOptions(initial_delay=2))

        )
    )